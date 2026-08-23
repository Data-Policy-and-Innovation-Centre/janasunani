"""Create the working-draft, non-technical Janasunani 2.0 officer brief.

This generator summarizes currently available evidence. It does not certify
the full-benchmark publication gate in docs/value-add-report/README.md.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from dpic.branding import colors as brand_colors

from janasunani.evaluation.value_add_benchmark_facts import (
    DEFAULT_BUNDLE,
    BenchmarkFacts,
    category_benchmark_summary,
    load_benchmark_facts,
)


def _word_color(value: str) -> str:
    return value.removeprefix("#")


NAVY = _word_color(brand_colors.PRIMARY)
TEAL = _word_color(brand_colors.BLUE)
GOLD = _word_color(brand_colors.ORANGE)
PALE_TEAL = _word_color(brand_colors.CARD_FILL_ALT)
PALE_GOLD = _word_color(brand_colors.CARD_FILL)
PALE_BLUE = _word_color(brand_colors.LIGHT_GREYSTONE)
ALT_FILL = _word_color(brand_colors.BACKGROUND)
MID_GREY = _word_color(brand_colors.TEXT_BODY)
LIGHT_GREY = _word_color(brand_colors.BORDER)
WHITE = _word_color(brand_colors.WHITE)
CONTENT_WIDTH_DXA = 10035
TABLE_INDENT_DXA = 140

DEFAULT_OUTPUT = Path(
    "docs/value-add-report/Janasunani_2.0_IAS_Officer_Brief_August_2026.docx"
)


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(cell, *, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_table_geometry(
    table,
    widths: tuple[int, ...],
    *,
    indent_dxa: int = TABLE_INDENT_DXA,
) -> None:
    """Apply fixed A4 brief geometry to a table and all of its cells.

    Named ``standard_business_brief`` overrides: A4 paper, 1.65 cm side
    margins, Aptos typography and the DPIC navy/teal/gold palette. Widths are
    DXA and must fill the 10,035-DXA content area exactly.
    """

    if len(widths) != len(table.columns) or sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must match columns and fill the content area")
    table.autofit = False
    properties = table._tbl.tblPr
    for tag, value in (
        ("w:tblW", CONTENT_WIDTH_DXA),
        ("w:tblInd", indent_dxa),
    ):
        node = properties.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            properties.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Twips(width)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, *, size: float = 10, bold: bool = False, color: str = NAVY) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _rule(paragraph, color: str = TEAL, size: int = 8) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _section_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Brief Section Label")
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(5)
    _rule(paragraph)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEAL)


def _title(document: Document, text: str, *, size: int = 26) -> None:
    paragraph = document.add_paragraph(style="Brief Title")
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def _body(document: Document, text: str, *, size: float = 11, color: str = MID_GREY, after: int = 7) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _callout(document: Document, heading: str, text: str, *, fill: str = PALE_GOLD) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    _shade(cell, fill)
    _cell_margins(cell, top=160, bottom=160, start=180, end=180)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(heading)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(GOLD if fill == PALE_GOLD else TEAL)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    _set_table_geometry(table, (CONTENT_WIDTH_DXA,), indent_dxa=180)


def _headline_cards(document: Document, facts: BenchmarkFacts) -> None:
    text = facts.latency["input_paths"]["text"]["e2e"]
    pdf = facts.latency["input_paths"]["document"]["e2e"]
    routing_top3 = facts.routing_all["top_k_accuracy"]["3"]
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    values = (
        (
            f"{text['p50']:.2f} seconds",
            f"Warm typed grievance p50; PDF p50 {pdf['p50']:.2f} s",
            PALE_TEAL,
            TEAL,
        ),
        (
            f"{routing_top3:.0%}",
            "Historical destination appears in the top 3 suggestions",
            PALE_BLUE,
            NAVY,
        ),
        ("5.07 filings", "Per inferred problem in the Sambalpur 2024 review slice", PALE_GOLD, GOLD),
    )
    for cell, (value, label, fill, color) in zip(table.rows[0].cells, values, strict=True):
        _shade(cell, fill)
        _cell_margins(cell, top=180, bottom=180, start=160, end=160)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(value)
        r.bold = True
        r.font.name = "Aptos Display"
        r.font.size = Pt(22)
        r.font.color.rgb = RGBColor.from_string(color)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.name = "Aptos"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(NAVY)
    _set_table_geometry(table, (3345, 3345, 3345), indent_dxa=160)


def _three_steps(document: Document) -> None:
    table = document.add_table(rows=3, cols=3)
    table.autofit = False
    rows = (
        ("1", "Read once", "The pilot is designed to bring text and useful pages into one officer-ready packet."),
        ("2", "Decide with context", "Advisory low-signal, category, duplicate and route suggestions are tested without automatic rejection."),
        ("3", "Supervise the queue", "The review view separates filings, inferred problems and signatories; live panel integration remains to be verified."),
    )
    for row, (number, heading, text) in zip(table.rows, rows, strict=True):
        _shade(row.cells[0], TEAL)
        _set_cell_text(row.cells[0], number, size=16, bold=True, color=WHITE)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(row.cells[1], PALE_TEAL)
        _set_cell_text(row.cells[1], heading, size=11, bold=True)
        _shade(row.cells[2], "F7F9FB")
        _set_cell_text(row.cells[2], text, size=10.5)
    _set_table_geometry(table, (800, 2100, 7135))


def _evidence_table(document: Document, facts: BenchmarkFacts) -> None:
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    headings = ("What is measured now", "What it means", "What it does not yet prove")
    for cell, heading in zip(table.rows[0].cells, headings, strict=True):
        _shade(cell, NAVY)
        _set_cell_text(cell, heading, size=9.5, bold=True, color=WHITE)
    _set_repeat_table_header(table.rows[0])
    route = facts.routing_all
    route_info = facts.routing_informative
    category = facts.categorization
    summary = facts.summary
    weak_count = facts.weak_labels["eligible_ticket_labels"]["valid_single_label"]
    rows = (
        (
            f"Routing: {route['accuracy']:.2%} top-1 and "
            f"{route['top_k_accuracy']['3']:.2%} top-3 on {route['n']:,} later "
            f"cases; {route_info['accuracy']:.2%} / "
            f"{route_info['top_k_accuracy']['3']:.2%} where intake category is informative.",
            "Moving from one suggestion to three raises historical-destination coverage; officer usefulness still has to be measured.",
            "Historical agreement is not proof that the destination was legally correct or produced the best outcome.",
        ),
        (
            f"Category shortlist: {category['accuracy']:.2%} top-1 and "
            f"{category['top_k_accuracy']['3']:.2%} top-3 historical-label "
            f"agreement on a viewed 2024 test of {category['n']:,} later cases.",
            "The local candidate can often place the recorded category in a short "
            "officer-review list.",
            "The split prevents exact-text leakage, but the result is not policy "
            "correctness or release evidence and must not auto-assign a case.",
        ),
        (
            f"Summary baseline: {summary['critical_fact_recall']['successes']}/"
            f"{summary['critical_fact_recall']['n']} critical facts retained; "
            f"{summary['usable_without_edit_rate']['successes']}/"
            f"{summary['generated_n']} drafts usable without edit.",
            "The local model was conservative about invention: the single judge found no unsupported or contradictory generated case.",
            f"It repeated residual identifying detail in "
            f"{summary['pii_leak_case_rate']['successes']}/{summary['generated_n']} "
            "drafts and failed the vague-input/Odia skip cases; officer validation is still required.",
        ),
        (
            f"Low-signal taxonomy: {weak_count:,} non-conflicting administrative "
            "weak labels across underspecified, irrelevant, outside-purview and policy cases.",
            "Different queue problems can be handled differently instead of calling every difficult filing spam.",
            "These are train-only weak labels. Accuracy, false-positive rate and a production threshold need officer-adjudicated gold.",
        ),
        (
            "Dedup review slice: 55,544 filings grouped into 10,963 inferred problems and 8,560 signatories.",
            "The officer can distinguish one repeat filer from a broad citizen campaign and preserve the campaign signal.",
            "The groups are reviewable operational evidence, not adjudicated duplicate precision or recall.",
        ),
        (
            "Sarvam: 56 cached paired page successes from an interrupted run; normalized outputs differed on every pair.",
            "The provider path, checkpointing and benchmark evidence are wired without spending new credits.",
            "Without hand transcription, divergence is not OCR accuracy and does not show which output is better.",
        ),
    )
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = PALE_BLUE if index % 2 == 0 else ALT_FILL
        for cell, value in zip(cells, values, strict=True):
            _shade(cell, fill)
            _set_cell_text(cell, value, size=9.5)
    _set_table_geometry(table, (3200, 3000, 3835))


def _decision_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    for cell, heading in zip(
        table.rows[0].cells,
        ("Decision", "Immediate output", "Success measure", "Safety condition"),
        strict=True,
    ):
        _shade(cell, NAVY)
        _set_cell_text(cell, heading, size=9.5, bold=True, color=WHITE)
    _set_repeat_table_header(table.rows[0])
    rows = (
        (
            "Approve a small officer-adjudicated gold set",
            "Reliable low-signal/actionability, summary and correct-authority scorecards",
            "Per-class precision/recall, harmful-review rate, summary factuality",
            "Raw citizen text stays governed; weak labels remain train-only",
        ),
        (
            "Run the advisory system in shadow mode",
            "Predictions recorded but hidden from officers",
            "Availability, latency, fallback, slice stability",
            "No automatic discard or assignment",
        ),
        (
            "Lock and run a bounded stepped rollout",
            "A fair comparison across pre-treatment office clusters",
            "7-day first action, transfer-free assignment, 90-day resolution",
            "ITT analysis, censoring, spillover checks and pause rules locked first",
        ),
    )
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = PALE_TEAL if index == 0 else (PALE_BLUE if index == 1 else PALE_GOLD)
        for cell, value in zip(cells, values, strict=True):
            _shade(cell, fill)
            _set_cell_text(cell, value, size=9.2)
    _set_table_geometry(table, (2200, 2400, 2500, 2935))


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(NAVY)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    for name in ("Brief Title", "Brief Section Label"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    title_style = styles["Brief Title"]
    title_style.font.name = "Aptos Display"
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string(NAVY)
    title_style.paragraph_format.space_before = Pt(3)
    title_style.paragraph_format.space_after = Pt(7)
    label_style = styles["Brief Section Label"]
    label_style.font.name = "Aptos Display"
    label_style.font.size = Pt(10)
    label_style.font.bold = True
    label_style.font.color.rgb = RGBColor.from_string(TEAL)
    label_style.paragraph_format.space_before = Pt(4)
    label_style.paragraph_format.space_after = Pt(5)

    document.core_properties.author = "Data, Policy and Innovation Centre"
    document.core_properties.subject = "Janasunani 2.0 evidence and decision support"

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("GOVERNMENT OF ODISHA  ×  UNIVERSITY OF CHICAGO TRUST")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    _rule(paragraph, TEAL, 5)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Janasunani 2.0  •  Officer decision brief  •  Working draft, August 2026")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)


def create_brief(destination: Path, *, benchmark_bundle: Path = DEFAULT_BUNDLE) -> None:
    facts = load_benchmark_facts(benchmark_bundle)
    document = Document()
    _configure(document)

    _section_label(document, "Decision brief | four-minute read")
    _title(document, "Janasunani 2.0: what changes for an officer—and why it matters", size=29)
    _body(
        document,
        "The proposed change is simple: after end-to-end verification, a grievance can arrive as an officer-ready packet with useful text, measured redaction, an advisory category and route shortlist, and a warning when the filing needs clarification. Redaction has known misses, and the officer remains the decision-maker.",
        size=12,
        color=NAVY,
        after=12,
    )
    _headline_cards(document, facts)
    _body(
        document,
        "These numbers measure three different things: technical speed, historical routing agreement, and duplicate-adjusted workload. They should not be added together or described as time saved.",
        size=9.5,
        after=10,
    )
    _callout(
        document,
        "The value proposition",
        "Test whether an officer-ready packet reduces reading; avoid forced guesses when a complaint is unclear; offer a short route list instead of one brittle answer; and distinguish repeat filing from broad citizen demand.",
        fill=PALE_TEAL,
    )
    _body(
        document,
        "Scale of the administrative problem: 688,301 complaints were filed in 2024–25, and roughly three in four filings carry at least one attachment. Among forwarded cases, the observed administrative interval to first forward was 1.7 days at the median. In controlled laptop tests the software prepared a warm typed filing in seconds; the full scanned browser path still needs verification, and only a controlled rollout can establish how much officer or citizen time it saves.",
        size=10.5,
        after=5,
    )

    document.add_page_break()
    _section_label(document, "What the officer sees")
    _title(document, "One screen, three decisions")
    _three_steps(document)
    _body(document, "The serving contract is advisory throughout: a flag must never block a filing, a category must never auto-assign a case, and a route suggestion must never replace the officer’s judgment.", size=10.5, after=10)
    _callout(
        document,
        "Example of better triage",
        "An abusive or content-free sentence is marked low-signal, its language is left unknown, and category and summary are skipped. The evaluation taxonomy keeps an incomplete request separate from irrelevant, outside-purview and policy-blocked cases. A checksummed binary model can now serve the advisory review decision, but it does not invent a five-class reason and is not approved for release. This fixes the screenshot failure without pretending all four cases are ‘spam’.",
    )
    _section_label(document, "What a supervisor gains")
    _body(
        document,
        "The Sambalpur review demonstrates workload as filings, inferred problems and distinct signatories; the live supervisor-panel integration remains to be verified. The distinction matters: a single repeat filer and a broad citizen campaign need very different responses even when both create large filing counts.",
        size=11,
        color=NAVY,
    )
    _callout(
        document,
        "What remains local",
        "The production contract requires models to load from pinned local release bytes or a DVC mirror; no reviewed production manifest is active yet. Serving does not follow a moving MLflow alias or download a public model at startup. Sarvam remains an explicit, auditable external route with a local fallback.",
        fill=PALE_BLUE,
    )

    document.add_page_break()
    _section_label(document, "The evidence, in plain language")
    _title(document, "What the current numbers support—and what they do not")
    _evidence_table(document, facts)
    _body(
        document,
        f"Four practical findings stand out. First, the tracked frontier-adjudicated "
        f"binary development test caught {facts.actionability['confusion']['true_review']}/"
        f"{facts.actionability['actual_review']} complaints needing extra review, while "
        f"also sending {facts.actionability['confusion']['false_review']}/"
        f"{facts.actionability['confusion']['true_actionable'] + facts.actionability['confusion']['false_review']} "
        f"ordinary complaints to review. Its checksummed binary artifact can serve "
        f"that advisory review decision, but does not assign five-class reasons and "
        f"is not release-eligible. Second, a separate chronological category model "
        f"measured {category_benchmark_summary(facts.categorization)}; this "
        f"is historical-label agreement, not policy correctness. Third, the local "
        f"BART baseline retained {facts.summary['critical_fact_recall']['successes']}/"
        f"{facts.summary['critical_fact_recall']['n']} critical facts and produced "
        f"{facts.summary['usable_without_edit_rate']['successes']}/"
        f"{facts.summary['generated_n']} drafts usable without edit; privacy and "
        f"skip failures make it a repair baseline. Fourth, category plus district "
        f"places the later historical "
        f"destination in the top three for {facts.routing_all['top_k_accuracy']['3']:.2%} "
        f"of cases, rising to {facts.routing_informative['top_k_accuracy']['3']:.2%} "
        "where the intake category is informative. All need a newly frozen, "
        "officer-reviewed release set.",
        size=10,
        after=8,
    )
    _callout(
        document,
        "Honest claim",
        f"The tracked development run completed "
        f"{facts.latency['completed_attempts']}/{facts.latency['attempts']} synthetic "
        f"attempts with {facts.latency['failed_attempts']} failures. Impact remains "
        f"unmeasured ({facts.impact_available_required}/{facts.impact_required} required "
        "impact artifacts available): no officer minutes saved, correct legal authority, "
        "faster resolution, or citizen satisfaction effect is claimed. Those claims "
        "require exposure logging, adjudication and a locked rollout.",
    )

    document.add_page_break()
    _section_label(document, "Decision requested")
    _title(document, "Turn a promising tool into measurable public value")
    _body(
        document,
        "The next investment is not another dashboard. It is a small amount of governed officer judgment plus a bounded rollout, so that model quality can be connected to fewer handoffs, earlier action and faster resolution.",
        size=11.5,
        color=NAVY,
    )
    _decision_table(document)
    _section_label(document, "The impact chain")
    _body(
        document,
        "Better model → officer sees and accepts a useful suggestion → fewer touches and handoffs → earlier meaningful action → better 30/90-day resolution, fewer repeat contacts and improved citizen-reported satisfaction.",
        size=13,
        color=TEAL,
        after=10,
    )
    _callout(
        document,
        "Recommended decision",
        "Approve creation and officer review of the release evidence set; begin shadow deployment only after the real browser/model path passes preflight. Lock the pilot and fixed-horizon satisfaction invitation rule before any treatment outcomes are viewed, and keep auto-reject and auto-assignment at zero until the agreed safety and quality gates pass.",
        fill=PALE_GOLD,
    )
    _body(
        document,
        f"Source note: benchmark-backed figures come from full bundle "
        f"{facts.bundle_id}. Dedup and descriptive administrative findings retain their "
        "separate cited sources. Full definitions and reproduction limits remain in the "
        "long report and docs/QUALITY_BENCHMARKS.md.",
        size=8.5,
        after=0,
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--benchmark-bundle", type=Path, default=DEFAULT_BUNDLE)
    args = parser.parse_args()
    create_brief(args.output, benchmark_bundle=args.benchmark_bundle)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
