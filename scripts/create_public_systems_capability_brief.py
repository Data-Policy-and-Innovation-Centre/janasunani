"""Create the working-draft brief for prospective public-system partners.

This generator summarizes currently available evidence. It does not certify
the full-benchmark publication gate in docs/value-add-report/README.md.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from create_officer_brief import (
    GOLD,
    NAVY,
    PALE_BLUE,
    PALE_GOLD,
    PALE_TEAL,
    TEAL,
    WHITE,
    _body,
    _callout,
    _cell_margins,
    _configure,
    _section_label,
    _set_cell_text,
    _set_repeat_table_header,
    _set_table_geometry,
    _shade,
    _title,
)
from janasunani.evaluation.value_add_benchmark_facts import (
    DEFAULT_BUNDLE,
    BenchmarkFacts,
    category_benchmark_summary,
    load_benchmark_facts,
)


DEFAULT_OUTPUT = Path(
    "docs/value-add-report/Janasunani_2.0_Public_Systems_Capability_Brief_August_2026.docx"
)


def _cards(document: Document, facts: BenchmarkFacts) -> None:
    table = document.add_table(rows=1, cols=3)
    values = (
        (
            "Local-first",
            "Production models can run on controlled infrastructure, with external providers optional and auditable.",
            PALE_TEAL,
            TEAL,
        ),
        (
            f"{facts.routing_all['top_k_accuracy']['3']:.0%}",
            "In Janasunani, the later historical destination appears in the top three suggestions overall.",
            PALE_BLUE,
            NAVY,
        ),
        (
            "4 questions",
            "Is it actionable? What is it about? Where should it go? Is it one filing or a wider problem?",
            PALE_GOLD,
            GOLD,
        ),
    )
    for cell, (value, label, fill, color) in zip(table.rows[0].cells, values, strict=True):
        _shade(cell, fill)
        _cell_margins(cell, top=180, bottom=180, start=160, end=160)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(value)
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(21)
        run.font.color.rgb = RGBColor.from_string(color)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        run.font.name = "Aptos"
        run.font.size = Pt(9.2)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    _set_table_geometry(table, (3345, 3345, 3345), indent_dxa=160)


def _value_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    for cell, heading in zip(
        table.rows[0].cells,
        ("Common public-system problem", "What our approach adds", "Value to measure"),
        strict=True,
    ):
        _shade(cell, NAVY)
        _set_cell_text(cell, heading, size=9.5, bold=True, color=WHITE)
    _set_repeat_table_header(table.rows[0])
    rows = (
        (
            "Case files arrive as short text, scans and mixed attachments.",
            "Governed extraction, measured redaction, page filtering and an officer-readable packet.",
            "Reading time, first meaningful action, model availability.",
        ),
        (
            "Incomplete, irrelevant and outside-purview requests are mixed into one queue.",
            "Separate actionability reasons and safe abstention; never a blanket ‘spam’ rejection.",
            "Clarification loops, review precision, actionable cases wrongly flagged.",
        ),
        (
            "Routing depends on memory and inconsistent labels.",
            "A ranked shortlist blending local text models, history, geography and rules.",
            "Transfer-free first assignment, handoffs, time to the responsible office.",
        ),
        (
            "Dashboards count filings but cannot distinguish repetition from broad demand.",
            "Access-controlled grouping over redacted text, with residual privacy risk governed locally.",
            "Reviewable extra matches, false merges, issues and citizens—not just filings.",
        ),
    )
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = PALE_BLUE if index % 2 == 0 else "F8FAFC"
        for cell, value in zip(cells, values, strict=True):
            _shade(cell, fill)
            _set_cell_text(cell, value, size=9.5)
    _set_table_geometry(table, (2900, 3600, 3535))


def _approach_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    for cell, heading in zip(
        table.rows[0].cells,
        ("Layer", "Reusable capability", "Deliverable"),
        strict=True,
    ):
        _shade(cell, NAVY)
        _set_cell_text(cell, heading, size=9.5, bold=True, color=WHITE)
    _set_repeat_table_header(table.rows[0])
    rows = (
        ("1. Establish truth", "Reconcile records, define denominators and freeze a baseline.", "A metric registry and reproducible scorecard."),
        ("2. Protect data", "Redact before downstream analysis; declare each trust boundary.", "A governed local dataset and egress controls."),
        ("3. Build decision support", "Start with cheap local baselines; compare pretrained candidates on identical cases.", "Versioned actionability, category, summary and route candidates."),
        ("4. Serve safely", "Pin exact model bytes and parameters; abstain and fall back when evidence is weak.", "A release manifest with local rollback."),
        ("5. Prove value", "Log what was shown and what the officer did; connect model quality to workflow and citizen outcomes.", "A shadow run followed by a pre-agreed controlled pilot."),
    )
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = PALE_TEAL if index in {0, 4} else (PALE_BLUE if index % 2 else "F8FAFC")
        for cell, value in zip(cells, values, strict=True):
            _shade(cell, fill)
            _set_cell_text(cell, value, size=9.3)
    _set_table_geometry(table, (1600, 4100, 4335))


def _engagement_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=4)
    for cell, heading in zip(
        table.rows[0].cells,
        ("Phase", "We do", "Partner receives", "Go / no-go evidence"),
        strict=True,
    ):
        _shade(cell, NAVY)
        _set_cell_text(cell, heading, size=9.2, bold=True, color=WHITE)
    _set_repeat_table_header(table.rows[0])
    rows = (
        (
            "Value diagnostic",
            "Reconcile the case flow, sample failure modes and establish baselines.",
            "A decision memo, metric registry and prioritized use cases.",
            "Is there enough reliable data and operational headroom?",
        ),
        (
            "Shadow pilot",
            "Run local candidate models without changing officer or citizen outcomes.",
            "Quality, coverage, latency, failure and subgroup scorecards.",
            "Does the tool clear safety and usefulness gates?",
        ),
        (
            "Measured rollout",
            "Expose advisory outputs in stages and record officer decisions safely.",
            "Evidence on touches, handoffs, first action and 30/90-day outcomes.",
            "Does it improve service without creating new harm?",
        ),
    )
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = PALE_TEAL if index == 0 else (PALE_BLUE if index == 1 else PALE_GOLD)
        for cell, value in zip(cells, values, strict=True):
            _shade(cell, fill)
            _set_cell_text(cell, value, size=9.1)
    _set_table_geometry(table, (1500, 2500, 2700, 3335))


def create_brief(destination: Path, *, benchmark_bundle: Path = DEFAULT_BUNDLE) -> None:
    facts = load_benchmark_facts(benchmark_bundle)
    document = Document()
    _configure(document)
    header_run = document.sections[0].header.paragraphs[0].runs[0]
    header_run.text = "DATA, POLICY AND INNOVATION CENTRE  •  PUBLIC SYSTEMS"
    footer_run = document.sections[0].footer.paragraphs[0].runs[0]
    footer_run.text = "DPIC  •  Public systems capability brief  •  Working draft, August 2026"

    _section_label(document, "Capability brief | for public-system partners")
    _title(document, "From case records to measurable public value", size=30)
    _body(
        document,
        "We help public agencies build governed decision support over unstructured complaints and case histories, then measure whether it makes service safer or faster—without making an expensive external model the permanent production dependency.",
        size=12.5,
        color=NAVY,
        after=12,
    )
    _cards(document, facts)
    _body(
        document,
        "Janasunani is the proof environment, not the limit of the approach. The architecture can be adapted wherever a public team must read mixed records, decide whether a request can be acted on, route it, find repetition and show whether service improved.",
        size=10.5,
        after=9,
    )
    _callout(
        document,
        "Where this transfers",
        "Grievance portals, welfare and benefit casework, municipal and utility complaints, helplines, regulatory case queues, inspection follow-up, and any service where free text must become a governed operational decision.",
        fill=PALE_TEAL,
    )
    _section_label(document, "The value we add")
    _value_table(document)

    document.add_page_break()
    _section_label(document, "A reusable delivery approach")
    _title(document, "Models are one layer; measurement and governance make them useful")
    _approach_table(document)
    _body(
        document,
        "We use separate frontier-model passes selectively to accelerate one-time adjudication and research, with explicit egress approval and known provenance limits. Production can remain local: classical text models, multilingual encoders and rules are benchmarked against the same frozen cases, then the cheapest model that clears the gate is pinned and served.",
        size=10.5,
        after=8,
    )
    _callout(
        document,
        "Proof point from Janasunani",
        f"A chronological benchmark places the historical destination in the top "
        f"three suggestions for {facts.routing_all['top_k_accuracy']['3']:.2%} of "
        f"{facts.routing_all['n']:,} cases, rising to "
        f"{facts.routing_informative['top_k_accuracy']['3']:.2%} where intake data "
        f"is informative. In a separate {facts.actionability['n']}-case "
        f"frontier-adjudicated binary development test, the validation-selected "
        f"local model caught {facts.actionability['confusion']['true_review']}/"
        f"{facts.actionability['actual_review']} complaints needing review while "
        f"also flagging {facts.actionability['confusion']['false_review']}/"
        f"{facts.actionability['confusion']['true_actionable'] + facts.actionability['confusion']['false_review']} "
        "ordinary complaints. Its checksummed binary artifact is serving-compatible "
        "for advisory review, but does not assign five-class reasons and is not "
        "release-eligible. A separate category benchmark measured "
        f"{category_benchmark_summary(facts.categorization)}. That is "
        "historical-label agreement, not policy correctness. "
        f"A separate local-BART summary baseline retained "
        f"{facts.summary['critical_fact_recall']['successes']}/"
        f"{facts.summary['critical_fact_recall']['n']} critical facts and produced "
        f"{facts.summary['usable_without_edit_rate']['successes']}/"
        f"{facts.summary['generated_n']} drafts usable without edit; residual-PII "
        "and skip failures keep it below release quality. "
        "These are proof points, not release or impact claims.",
        fill=PALE_BLUE,
    )
    _callout(
        document,
        "Why this is different from a generic AI demo",
        f"Every number keeps its denominator and limitation. The current full bundle "
        f"is not publication-ready and has {facts.impact_available_required}/"
        f"{facts.impact_required} required impact artifacts available. The release "
        "mechanism can pin parameters and support rollback once an approved manifest "
        "is activated. Low-confidence cases abstain. External egress is explicit. "
        "The evaluation follows the chain from model to officer behavior to workflow "
        "to citizen outcome.",
        fill=PALE_GOLD,
    )

    document.add_page_break()
    _section_label(document, "How a prospective partner can start")
    _title(document, "Begin with evidence; scale only after value is visible")
    _engagement_table(document)
    _section_label(document, "What we need from a partner")
    _body(
        document,
        "A named service problem; a governed extract or secure on-premise access; the officers who know what a good decision looks like; and agreement on the outcome that matters. Raw citizen text need not leave the partner’s controlled environment for production.",
        size=11,
        color=NAVY,
    )
    _section_label(document, "What the partner should expect")
    _body(
        document,
        "A working local pipeline, transparent scorecards, a model and parameter catalog, clear fallbacks, and a decision on whether the use case deserves a pilot. If the evidence is weak, the deliverable is an honest no-go and the data needed to revisit it—not an inflated accuracy slide.",
        size=11,
        color=NAVY,
    )
    _callout(
        document,
        "The proposition",
        "Use advanced models where they create learning; keep recurring production cost and data exposure low; and measure success in officer effort, handoffs, time to action, citizen resolution and satisfaction response—not merely an offline model score.",
        fill=PALE_TEAL,
    )
    _body(
        document,
        f"Evidence note: benchmark-backed figures come from development bundle "
        f"{facts.bundle_id}; publication_ready={str(facts.publication_ready).lower()}. "
        "A new partner begins with its own baseline, taxonomy, governed sample and "
        "release gates.",
        size=8.5,
        after=0,
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--benchmark-bundle", type=Path, default=DEFAULT_BUNDLE)
    args = parser.parse_args()
    create_brief(args.output, benchmark_bundle=args.benchmark_bundle)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
