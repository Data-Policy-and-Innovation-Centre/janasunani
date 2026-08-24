"""Apply evidence-status corrections to the August 2026 working-draft DOCX.

The source report was produced without a durable generator.  This narrow,
idempotent patcher keeps the existing layout and makes the evidence corrections
reproducible until the report is regenerated from a complete, versioned bundle
of quality, timing, and impact benchmark results. It does not certify that the
final publication gate in docs/value-add-report/README.md has passed.
"""

from __future__ import annotations

import argparse
from io import BytesIO
import os
from pathlib import Path
import tempfile
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from dpic.branding import colors as brand_colors

from docx_archive import canonicalize_docx_archive
from janasunani.evaluation.value_add_benchmark_facts import (
    DEFAULT_BUNDLE,
    BenchmarkFacts,
    category_benchmark_summary,
    load_benchmark_facts,
)


DEFAULT_REPORT = Path(
    "docs/value-add-report/Janasunani_2.0_Value_Add_Report_August_2026.docx"
)
DEFAULT_SOURCE = Path(
    "docs/value-add-report/templates/Janasunani_2.0_Value_Add_Report_August_2026.source.docx"
)

FIG_PRIMARY = brand_colors.PRIMARY
FIG_BLUE = brand_colors.BLUE
FIG_ACCENT = brand_colors.ORANGE
FIG_TEXT = brand_colors.TEXT_BODY
FIG_TEXT_SECONDARY = brand_colors.TEXT_SECONDARY
FIG_BORDER = brand_colors.BORDER
FIG_BACKGROUND = brand_colors.BACKGROUND
FIG_CARD = brand_colors.CARD_FILL
FIG_CARD_ALT = brand_colors.CARD_FILL_ALT


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is not None:
        return
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


PARAGRAPH_REPLACEMENTS = {
    "1   How grievances work today  — the status quo in numbers\t3": "1   How grievances work today  — the status quo in numbers\t2",
    "2   What we built  — the same workflow, now machine-readable\t4": "2   What we built  — the same workflow, now machine-readable\t3",
    "3   Faster and more reliable  — document processing before vs after\t5": "3   Faster and more reliable  — document processing before vs after\t4",
    "4   Keeping the queue clean  — spam and duplicates\t8": "4   Keeping the queue clean  — spam and duplicates\t6",
    "5   What are people complaining about  — hotspots, spikes, and workload\t10": "5   What are people complaining about  — hotspots, spikes, and workload\t7",
    "6   Who is complaining  — geography, gender, and channels\t13": "6   Who is complaining  — geography, gender, and channels\t9",
    "7   What officers see day-to-day  — real-time help while a case is open\t15": "7   What officers see day-to-day  — real-time help while a case is open\t10",
    "8   Safeguards and honesty  — what we don’t claim yet\t16": "8   Safeguards and honesty  — what we don’t claim yet\t11",
    "A   Annex — metric registry, sources, and how to reproduce\t17": "A   Annex — metric registry, sources, and how to reproduce\t12",
    "Janasunani receives about 1.4 million grievances a year. Three in four arrive with at least one attached document — together more than a million files. Today, an officer must open each filing, read the grievance and its attachments, decide what the complaint is about, and route it to the office that can act. That first routing step — understanding the complaint well enough to send it to the right place — takes a median of 1.7 days among forwarded cases, and a mean of 20.8 days because a long tail of cases takes much longer. Email filings take 17.7 days at the median; the Collector’s office, which handles one in seven routed complaints, takes 4 days at the median and 30 days at the mean.": "Between July 2021 and June 2025 Janasunani received 1,371,288 complaints; 688,301 arrived in 2024–25. Three in four filings carry at least one attachment, so officers must read the grievance and its documents before deciding what it is about and where it should go. Among forwarded cases, the administrative interval from receipt to first forward is 1.7 days at the median and 20.8 days at the mean; email is 17.7 days at the median, while the Collector’s office is 4 days at the median and 30 days at the mean.",
    "1. Faster first response — from reading to reviewing. The system reads the grievance and its scanned pages, removes personal information, labels the useful pages, summarises what matters, suggests a category, and proposes where to route it — in seconds for typed text, minutes for scans — so the officer starts from a first draft, not a blank page. Officers keep full authority: they accept, edit, or reject every suggestion.": "1. Faster first response — from reading to reviewing. The intended workflow extracts and redacts text, filters pages, and prepares advisory category, summary and route suggestions. Warm typed requests have completed in seconds in controlled laptop tests; the full scanned browser/model path and officer accept/edit/reject events are not yet verified. Redaction has measured misses, and officers retain full authority.",
    "3. Intelligence that earns a decision — hotspots with denominators, spikes explained, and workload counted as problems not filings. Geography is shown as complaints per 1,000 residents; time is shown against last year, not last month; every spike carries three numbers (filings / distinct problems / distinct citizens) so a campaign and a wave of unrelated problems — identical in a raw count — trigger different responses. Block-level diagnostics surface the blocks that a district average would hide.": "3. Intelligence that earns a decision — geographic counts with explicit limits, spikes explained, and workload counted as problems not filings. The first map release will show privacy-safe filing counts, inferred problems and distinct signatories; it will not publish per-capita or per-eligible rates until a current, documented denominator is governed. Time is compared with an appropriate historical baseline, and every spike carries three numbers (filings / inferred problems / distinct signatories) so a campaign and a diffuse wave — identical in a raw count — prompt different review.",
    "We rebuilt the grievance workflow as a sequence of governed, measurable steps. A raw filing — typed text or a scanned document — is extracted, redacted, triaged, classified, summarised, and routed, with a supervisor intelligence layer reading the same redacted corpus in parallel. Data stays on DPIC-controlled infrastructure; any external call (Sarvam) is declared, audited, and revocable — with a local fallback.": "We rebuilt the grievance workflow as a sequence of governed, measurable steps. The intended path extracts, redacts, triages, classifies, summarises and routes a filing, with a supervisor intelligence layer over governed redacted text. Most production inference is designed to remain on DPIC-controlled infrastructure. Sarvam is declared, audited and revocable with a local fallback; separately, a privacy-screened redacted sample was sent through hosted Codex for one-time development adjudication, not production.",
    "Figure 1 — The seven functional steps. Steps 1–6 are per-grievance (live in seconds/minutes); step 7 is corpus-level (nightly aggregates over the lake). Triage is advisory — it never changes submission status.": "Figure 1 — The seven functional steps. Steps 1–6 are implemented per-grievance stages, but the full scanned browser/model path is not yet verified; step 7 is corpus-level and some dedup-backed views still require integration. The triage contract is advisory and must never change submission status.",
    "The DSI Clinic built the first open-source pipeline on a 100,000-complaint sample (2025) — English-only, five stages, on an A100 GPU. We refolded it into a six-stage production pipeline (plus triage) that runs on DPIC infrastructure (CPU box + GPU box + laptop), replaced the unrecoverable PII model with a Presidio rebuild, added quality guards, and wired it to the live API. The comparison below is honest about what is like-for-like and what is not: baselines are historical reference, not thresholds.": "The DSI Clinic built the first open-source pipeline on a 100,000-complaint sample (2025) — English-only, five stages, on an A100 GPU. We implemented six local stages plus triage, replaced the unrecoverable PII model with a Presidio rebuild, and added quality guards and typed API contracts. This is not yet a verified production deployment: one-pass integration, live interface wiring and AWS activation remain outstanding. Baselines below are historical reference, not thresholds.",
    "Figure 2 — Time from filing to an officer-ready packet. Manual median 1.7 days is among forwarded complaints (n≈30k); live text median 4.4 s is n=8 warm submits on a laptop (PERFORMANCE.md). Log scale is required because the improvement is two orders of magnitude. First document after boot is slower (~9–10 s) while models warm; the summariser (~1.6 GB BART) is fetched once on first boot.": "Figure 2 — Two different clocks, shown side by side. The 1.7-day median is the observed administrative interval to first forward among ≈30,000 forwarded complaints; 4.4 seconds is technical pipeline latency from eight warm laptop submissions. This demonstrates processing speed, not days saved. Officer handling time and first-forward impact require exposure logging and a controlled rollout.",
    "Spam is not a block — it is a banner. A filing flagged as low-signal is still submitted; the officer sees “low-signal: <reason> (spam_score 0.82)” and decides. Prevalence is measured over redacted grievance text only (never raw grievance) and reported by district / category / mode / year — so a high-spam district is visible, not hidden in a state average. PPV / false-positive rate is measured against the two officer-confirmed spam-like families (details inadequate 39,964 + no specific grievance 16,375) on a deterministic 30% holdout by ticket hash; duplicate families are never counted as spam positives.": "Low-signal review is not a block. The evaluation taxonomy separates actionable, underspecified, irrelevant, outside-purview and policy-blocked cases; only advisory review or abstention is permitted and officers decide. Administrative templates provide 106,683 non-conflicting train-only weak labels, not adjudicated truth. The checksummed binary development artifact can now serve the advisory actionable-versus-review objective, but it does not produce the five-class reasons and is not release-eligible. Office variation fails the pooling gate (maximum total-variation distance 0.522), so no production threshold is approved until a stratified officer-adjudicated validation/test set exists.",
    "Duplicate-adjusted workload — the same three numbers, without a spike. The portal counts filings; the intelligence layer counts problems. For Sambalpur/2024: Filings 55,544 → Distinct problems 10,963 → Distinct citizens 8,560. That is the “true workload” for the slice. Both workload and spike share the same dedup_groups digest; the serving layer refuses a mixed snapshot (#137), so a stale index cannot silently undercount a surge.": "Duplicate-adjusted workload — the same three numbers, without a spike. For Sambalpur/2024: 55,544 filings → 10,963 inferred problems → 8,560 distinct signatories: 5.07 filings per inferred problem, or 80.3% fewer problem-units than filings. This is a reviewable operational view, not ground-truth workload. Workload and spike share the same dedup digest, and serving refuses a mixed snapshot (#137).",
    "Hotspot monitoring — per-1,000, not raw counts": "Geographic monitoring — governed counts first, rates later",
    "Headline complaint counts are not complaint rates — without an eligibility denominator (eligible households for PMAY, ration-card holders for food security, etc.), a district with 200 complaints looks the same whether it has 1,000 or 10,000 eligible households (2% vs 20%). The interim control is complaints per 1,000 general population (2021 projected district population, MoHFW 2020), excluding discards, tracked monthly.": "Headline filing counts are not complaint rates. Until a current, documented population or programme-eligibility denominator is governed, the first release will show privacy-safe filing counts, inferred-problem counts and distinct-signatory counts without ranking districts by incidence. Comparisons over time use each geography’s own historical baseline and always state the period and denominator.",
    "Month-to-month comparison is the operational use: a district whose per-1,000 rises faster than the state average is flagged — whether or not its absolute count is the highest. The roll-up goes district → block next (see below).": "The operational use is change detection: flag a district or block whose filing count rises unusually relative to its own history, then inspect whether the change reflects many signatories, repeat filing or inferred problem clusters. Cross-district rate comparisons remain disabled until a governed denominator exists; district-to-block drilldown follows only after the boundary crosswalk is deterministic.",
    "1. Eligibility denominators — link Janasunani to PMAY / pension / ration-card lists by district and block (per-eligible rates replace per-1,000). Highest priority.": "1. Eligibility denominators — later, link Janasunani to current, governed PMAY / pension / ration-card lists by district and block before publishing any per-eligible rate. The first release uses counts only. Highest priority.",
    "The system is human-in-the-loop. AI suggests; officers decide. Every automated suggestion is logged with who saw it, what model version was used, and whether the officer accepted, edited, or rejected it — that exposure log is what makes the future A/B evaluation possible. The live triage is advisory — it never auto-discards (even a “discard” recommendation in the old system is rendered as a flag, not a state change).": "The system is human-in-the-loop: AI suggests and officers decide; live triage never auto-discards. The model/release manifest can now pin exact versions locally, but the append-only exposure and later officer-decision events are still to be implemented. Until they exist, the report cannot claim acceptance, edit, override, time saved, or causal impact.",
    "Timing the officer feels. Cold start to /health: 19.4 s (models on disk) — the slow part is first boot when BART (~1.6 GB) is fetched from the hub (pre-warm the night before). Warm text grievance: median 4.44 s, mean 4.77 s (n=8, laptop). First request after boot ~9.5 s. Typed-text needs no tesseract/poppler at submit time beyond preflight checks; document upload renders via Poppler (pdftoppm/pdfinfo) then Tesseract (Oriya needs ori traineddata) before the same downstream steps. Every response includes extraction source (text vs ocr_model), spans (start/end over original text), and advisory fields — all typed and validated against the frozen frontend contract (janasunani/serving/schemas.py).": "Technical timing. Cold start to /health was 19.4 s with models already on disk. Serving now requires locally materialized, pinned BART bytes and performs no model download at startup. A warm typed grievance took a 4.44 s median and 4.77 s mean (n=8, laptop); the first request after boot was about 9.5 s. Document uploads render via Poppler and then use local OCR before the same downstream stages. These are technical timings, not officer time saved. Responses retain extraction source, original-text spans, and advisory fields under the typed serving contract.",
    "Evaluation discipline: a harness that measures and prints (evaluation/) is separate from a gate that fails a run when the number is bad (pipeline/pii_eval.py). DSI baselines are labelled reference_only=True (dsi_baselines.py) and the report renderer never colours them as targets. The lake is not PII-free — and we do not pretend it is (ROADMAP §3.2).": "Evaluation discipline: a harness that measures and prints (evaluation/) is separate from a gate that fails a run when a frozen number is bad (pipeline/pii_eval.py). The impact ladder is model quality → officer behavior → workflow outcome → citizen outcome. Current data can support selected model and workflow descriptions; officer behavior needs exposure/decision logging, correct authority needs adjudication, and causal citizen benefit needs a locked pilot. The lake is not PII-free — and we do not pretend it is (ROADMAP §3.2).",
    "The intelligence layer is not only a monthly report — it is a banner the officer sees while deciding, and a supervisor panel that refreshes from the same governed marts that produce the monthly findings. Both are built for the reality that a count without a denominator, a spike without a cause, and a queue with duplicates are worse than no information.": "The intelligence layer currently provides governed findings and a reviewable dedup slice. The officer banner and live supervisor panel are intended delivery surfaces, but their real-data integration is not yet verified and some spike/workload marts still await the dedup-index join. A count without a denominator, a spike without a cause and a queue with duplicates remain the design problem.",
    "Figure 7 — Left: learned crosswalk accuracy rises when district is added (≈12 pp from category-only to full). Right: Route 4’s median 48 days vs Route 2’s 23 days (PMAY, 2024–25); the 9-day Step 3 (BDO → Collector return) is the piloted reform target. Upper bound, not an estimate — Route 4 may be more complex on unobserved dimensions.": "Figure 7 — Left: the 60.9/67.5/72.8% crosswalk bars are historical in-sample resubstitution, retained only as upper-bound context. A chronological developmental holdout gives 45.14% top-1 and 69.04% top-3 for live category+district features (n=208,267), or 54.96%/79.68% on informative categories. Right: Route 4’s 48-day vs Route 2’s 23-day medians are descriptive and may reflect unobserved complexity; a pilot must measure any time effect.",
    "Plus the two engineering slices that are built and waiting for an overnight run and a key: dedup-index join into spike/workload marts (#78) and the Sarvam live comparison (needs SARVAM_API_KEY — few-hundred rupees on the paired 300-page sample; governance is recorded, cost is not a blocker). The A/B stepped-wedge design (AB_PLAN.md) is locked before any outcome data is viewed — so the August framework does not become a post-hoc story.": "Next engineering work is evidence-preserving rather than new provider spend: checkpoint each Sarvam page so an interruption cannot lose paid results, import the cached 56-page aggregate into the benchmark registry, and wire the dedup-index join into spike/workload marts (#78). The stepped-wedge A/B plan remains DRAFT; its unit map, estimands, extract hash, MDE and pause rules must be locked before any arm outcome is read.",
}

# The report was patched once before binary advisory serving was implemented.
# Keep that intermediate paragraph as an accepted source so the
# generator remains idempotent across both tracked states.
PARAGRAPH_ALTERNATES = {
    "Figure 7 — Left: learned crosswalk accuracy rises when district is added (≈12 pp from category-only to full). Right: Route 4’s median 48 days vs Route 2’s 23 days (PMAY, 2024–25); the 9-day Step 3 (BDO → Collector return) is the piloted reform target. Upper bound, not an estimate — Route 4 may be more complex on unobserved dimensions.": (
        "Figure 7 — Left: the 60.9/67.5/72.8% crosswalk bars are historical in-sample resubstitution, retained only as upper-bound context. A chronological developmental holdout gives 45.15% top-1 and 69.05% top-3 for live category+district features (n=208,267), or 54.96%/79.68% on informative categories. Right: Route 4’s 48-day vs Route 2’s 23-day medians are descriptive and may reflect unobserved complexity; a pilot must measure any time effect.",
    ),
    "Spam is not a block — it is a banner. A filing flagged as low-signal is still submitted; the officer sees “low-signal: <reason> (spam_score 0.82)” and decides. Prevalence is measured over redacted grievance text only (never raw grievance) and reported by district / category / mode / year — so a high-spam district is visible, not hidden in a state average. PPV / false-positive rate is measured against the two officer-confirmed spam-like families (details inadequate 39,964 + no specific grievance 16,375) on a deterministic 30% holdout by ticket hash; duplicate families are never counted as spam positives.": (
        "Low-signal review is not a block. The five classes are actionable, underspecified, irrelevant, outside purview, and policy-blocked; only an advisory review flag is permitted and officers decide. Administrative templates provide 106,683 non-conflicting train-only weak labels, not adjudicated truth. Office variation fails the pooling gate (maximum total-variation distance 0.522), so no PPV, false-positive rate, or production threshold is claimed until a stratified officer-adjudicated validation/test set exists.",
        "Low-signal review is not a block. The evaluation taxonomy separates actionable, underspecified, irrelevant, outside-purview and policy-blocked cases; only advisory review or abstention is permitted and officers decide. Administrative templates provide 106,683 non-conflicting train-only weak labels, not adjudicated truth. A separate binary development benchmark is not compatible with the five-class serving slot and produced no deployable artifact. Office variation fails the pooling gate (maximum total-variation distance 0.522), so no PPV, false-positive rate or production threshold is approved until a stratified officer-adjudicated validation/test set exists.",
    ),
}


CELL_REPLACEMENTS = {
    (
        2,
        0,
        0,
        1,
    ): "Median technical latency for a warm typed grievance (n=8). The 1.7-day first-forward interval is a different administrative clock; workflow time saved is not yet measured.",
    (
        3,
        0,
        0,
        0,
    ): "⚠  What we are not claiming. Technical latency is not officer time saved. Sarvam evidence is divergence/coverage, not OCR accuracy. Administrative discard templates are weak labels, not spam gold. Summary has only a small single-frontier-judge development baseline, with privacy and skip failures—not officer validation. Routing’s older 60.9/67.5/72.8% is in-sample, while the new chronological 2025 result is developmental because that test was viewed. Causal officer or citizen impact requires a locked pilot.",
    (
        9,
        5,
        2,
        0,
    ): "BART incumbent (facebook/bart-large-cnn revision 37f520fa…, ~1.6 GB). On a deterministic enriched 30-case redacted typed-text development set, a single frontier judge found 55/84 critical facts retained, 8/26 drafts usable without edit, no unsupported or contradictory generated case, and 4/26 residual-PII cases. Serving requires a pinned local release/DVC artifact.",
    (
        9,
        5,
        3,
        0,
    ): "The development baseline summarized all six cases the judge marked for skipping and skipped all four coherent Odia cases; post-summary privacy, better abstention and a newly frozen paired officer review are required before release.",
    (
        9,
        6,
        2,
        0,
    ): "MuRIL remains the serving incumbent. A separate local hashing candidate was evaluated on 2024 redacted typed text with chronological, exact-text-group-disjoint splits. Its viewed development test (n=3,160) reached 46.55% top-1, 90.89% top-3 and 36.49% macro-F1. This is historical-label agreement, not policy correctness or a release result.",
    (
        9,
        6,
        3,
        0,
    ): "Category suggestions can offer an officer a short list, but no automatic assignment is justified. The test was viewed, selective constraints were not met and language was unadjudicated. Promotion requires a newly frozen, officer-confirmed set with per-class, calibration, abstention and language/source slices.",
    (
        9,
        7,
        3,
        0,
    ): "The measured 4.44 seconds is technical latency, not officer handling time or time to first action. Those outcomes need exposure/decision logging and a controlled rollout.",
    (
        9,
        2,
        3,
        0,
    ): "About 77.9% of historical English pages passed three plausibility heuristics. Without hand transcription, that is not evidence that four in five pages were read correctly. Pages that fail the repetition guard are quarantined for review; OCR accuracy remains unmeasured.",
    (
        9,
        4,
        3,
        0,
    ): "Historical usefulness scores motivated skipping IDs and bills before summary generation. The new 30-case typed-text development baseline measures current BART failure modes but does not establish attachment prevalence, scan quality or officer usefulness; those claims need a newly frozen paired officer review.",
    (
        13,
        2,
        2,
        0,
    ): "55,544 / 10,963 = 5.07 filings per inferred problem, or 80.3% fewer problem-units than filings. These are reviewable inferred groups, not adjudicated ground truth.",
    (15, 0, 2, 0): "Why counts need context",
    (
        15,
        1,
        0,
        0,
    ): "Filing counts by district and block — descriptive workload only; not incidence or need.",
    (
        15,
        1,
        1,
        0,
    ): "Counts can locate workload and unusual changes within a geography, but they do not support fair rate comparisons across differently sized populations.",
    (
        15,
        1,
        2,
        0,
    ): "The first map release will not rank districts by per-capita or per-eligible rates until a current, documented denominator is governed.",
    (
        15,
        2,
        0,
        0,
    ): "Distinct signatories and inferred problems — reported beside filing counts.",
    (
        15,
        3,
        2,
        0,
    ): "A filing count alone cannot tell you whether to investigate broad service-delivery demand or a small set of repeat filers.",
    (
        15,
        4,
        0,
        0,
    ): "Rural Housing + PMAY filing volume — descriptive counts, with period and suppression status.",
    (
        18,
        0,
        0,
        0,
    ): "⚠  Eligibility denominators remain the highest-priority external data to acquire. A block with 200 housing complaints among 1,000 eligible households is a different fact from 200 among 10,000. Until current beneficiary lists are linked, documented and governed, the UI and report must not publish per-capita or per-eligible rates.",
    (
        19,
        5,
        1,
        0,
    ): "No hard rural/urban flag exists in the structured record. District and block fields can support privacy-safe raw aggregates after deterministic boundary reconciliation, but they do not identify rurality or incidence. Distinct-signatory and filings-per-signatory measures can describe repeat concentration without inventing a population denominator.",
    (
        20,
        0,
        0,
        0,
    ): "⚠  On rural/urban and disadvantaged groups — a candour note for reviewers. The complaint record does not contain a rural/urban flag or a caste/category field. The first map release may show privacy-safe district and block counts after deterministic boundary reconciliation, but it will not show per-capita rates, rural-versus-urban shares or SC/ST shares. Those claims require current, documented census, SECC or beneficiary denominators and approved linkage. We mark them as ‘needs linkage’ rather than filling the gap with a guess.",
    (
        21,
        1,
        1,
        0,
    ): "Advisory low-signal reason, never a rejection. The screenshot case now skips category/summary. On the canonical 57-case frontier-adjudicated binary development test, the local review candidate caught all 13 complaints needing extra review and sent 3 of 44 ordinary complaints to review. The checksummed binary artifact is serving-compatible for advisory review; it does not assign five-class reasons and is not release-eligible.",
    (
        21,
        1,
        2,
        0,
    ): "Administrative weak labels do not establish quality. The development result has no outside-purview support, wide intervals and a viewed test; an officer-adjudicated, stratified future set is still required. Officer always decides whether to proceed or seek clarification.",
    (
        21,
        3,
        1,
        0,
    ): "MuRIL incumbent with a ranked-category suggestion. A local hashing candidate reached 46.55% top-1 and 90.89% top-3 historical-label agreement on a viewed, exact-text-group-disjoint 2024 development test (n=3,160).",
    (
        21,
        3,
        2,
        0,
    ): "Does not auto-assign. The result is not policy correctness or release evidence; promotion requires a newly frozen, officer-confirmed set with per-class/top-k/calibration, abstention and language/source slices.",
    (
        21,
        4,
        1,
        0,
    ): "Local BART draft for English-compatible grievance text. Single-judge enriched development baseline: 65.48% critical-fact recall; 0/26 unsupported or contradictory cases; 8/26 usable without edit; 4/26 residual-PII cases.",
    (
        21,
        4,
        2,
        0,
    ): "The current guard passed all six judge-marked vague/underspecified cases to BART and skipped four coherent Odia cases. This is a repair baseline, not release or officer-validation evidence.",
    (
        21,
        5,
        2,
        0,
    ): "Learned means historical destination agreement, not correct authority or best outcome. Chronological developmental top-1 is 45.14% overall and 54.96% on informative categories; the older 60.9/67.5/72.8% figures are in-sample.",
    (
        23,
        2,
        1,
        0,
    ): "Citizen text remained on DPIC-controlled infrastructure, but the old runtime could still fetch public model bytes at startup. That made availability and exact rollback depend on mutable external state.",
    (
        23,
        2,
        2,
        0,
    ): "Trust tiers remain explicit. The serving contract resolves pinned local artifacts from an activated release manifest or DVC mirror and makes no registry/public-model call at startup; no reviewed production manifest is active yet. Sarvam traffic remains isolated to egress/ and kill-switch controlled. A shape-screened redacted sample was also sent to hosted Codex for one-time development adjudication; exact hidden prompts, sampling settings and provider-retention evidence were unavailable, so this is recorded as a limitation rather than production precedent.",
    (
        24,
        3,
        3,
        0,
    ): "Adjudicate the same grievance set by language and typed/scan source, freeze dataset and split fingerprints, then compare incumbent and candidate versions through the governed evaluation logger.",
    (
        24,
        4,
        1,
        0,
    ): "Officer-confirmed duplicate actions are a historical baseline. The Sambalpur/2024 index groups 55,544 filings into 10,963 inferred problems, but those inferred groups are not adjudicated recall or precision evidence.",
    (
        24,
        4,
        2,
        0,
    ): "No duplicate recall, candidate PPV, false-merge rate, or automation increment is claimed until the candidate pairs and clusters are adjudicated on a held-out sample.",
    (
        24,
        4,
        3,
        0,
    ): "Freeze an officer-held-out sample and publish recall, reviewable-candidate PPV, extra matches beyond the officer baseline, campaign preservation and singleton false merges with ticket/cluster bootstrap intervals.",
    (
        24,
        1,
        1,
        0,
    ): "Cached paired coverage/divergence only: a completed 5-page run and 56 paired successes from an interrupted 300-page run. Sarvam produced 1.3345× normalized characters on the 56-page aggregate; every normalized pair differed.",
    (
        24,
        1,
        2,
        0,
    ): "No OCR accuracy, handwriting or observed-language conclusion: there is no hand transcription. Nine attempted pages were excluded and credit exhaustion interrupted the larger run, so failures must be reported separately.",
    (
        24,
        3,
        1,
        0,
    ): "Category historical-label agreement: 46.55% top-1 / 90.89% top-3 / 36.49% macro-F1 on a viewed 2024 chronological, exact-text-group-disjoint development test (n=3,160). Not policy correctness or release evidence. Routing historical-destination agreement: 45.14% top-1 / 69.04% top-3 overall (n=208,267).",
    (
        24,
        5,
        1,
        0,
    ): "Historical destination agreement only. The live category+district developmental holdout is 45.14% top-1 (95% CI 44.94–45.36) and 69.04% top-3; process-time contrasts remain descriptive.",
    (
        24,
        5,
        2,
        0,
    ): "We do not report correct-authority rate or faster resolution caused by routing. Correctness needs jurisdiction adjudication; time/transfer benefit needs exposure logging and a locked rollout.",
    (
        24,
        5,
        3,
        0,
    ): "Lock the draft stepped-wedge plan before outcomes: immutable intake-office transfer-network clusters, ITT, censoring-aware 30/90-day endpoints, exposure/decision events, a fixed-horizon citizen-satisfaction invitation rule with response/missingness reporting, spillover sensitivity and pause rules.",
    (25, 10, 0, 0): "Routing — chronological developmental",
    (
        25,
        10,
        1,
        0,
    ): "Category+district: 45.14% top-1 (44.94–45.36), 69.04% top-3 / n=208,267; informative: 54.96% / 79.68% / n=142,181",
    (25, 10, 2, 0): "janasunani/evaluation/historical.py; docs/QUALITY_BENCHMARKS.md",
    (
        25,
        10,
        3,
        0,
    ): "janasunani-evaluate-routing (freeze a future slice before release)",
    (25, 9, 0, 0): "Categorization — chronological developmental",
    (
        25,
        9,
        1,
        0,
    ): "2024 exact-text-group-disjoint test n=3,160: 46.55% top-1 (44.82–48.29), 90.89% top-3, 95.19% top-5, 36.49% macro-F1; test viewed during development",
    (
        25,
        9,
        2,
        0,
    ): "outputs/evaluation/categorization_historical_v1.json; docs/QUALITY_BENCHMARKS.md",
    (
        25,
        9,
        3,
        0,
    ): "dvc repro --single-item categorization-historical-benchmark; freeze a new officer-confirmed test before release",
    (25, 12, 0, 0): "Actionability — development + weak-label audit",
    (
        25,
        12,
        1,
        0,
    ): "Canonical frontier-adjudicated test n=57: 94.74% accuracy; 13/13 review recall; 13/16 review precision; 3/44 actionable sent to review. Weak labels n=106,683; office max TV 0.522 (pooling gate fails).",
    (
        25,
        12,
        2,
        0,
    ): "docs/evidence/actionability_frontier_benchmark_reproducible.json; janasunani/evaluation/actionability.py; weak_labels.py",
    (
        25,
        12,
        3,
        0,
    ): "Serving-compatible advisory binary artifact; no five-class reasons or outside-purview support; not release-eligible; freeze officer-reviewed future test before promotion",
    (
        25,
        21,
        1,
        0,
    ): "DRAFT: stepped-wedge ITT framework; unit/event semantics, MDEs, extract hash and pause rules still must be locked",
    (
        25,
        22,
        1,
        0,
    ): "Cached: 5 completed pages + 56 paired successes from interrupted run; 7/127 accepted jobs failed; divergence/coverage only, no accuracy",
    (25, 22, 2, 0): "docs/evidence/sarvam_cached_benchmark.json; sarvam_scorecard.py",
    (
        25,
        22,
        3,
        0,
    ): "No paid rerun; add transcription/adjudication before a quality claim",
}


def _set_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _figure_png(fig) -> bytes:
    payload = BytesIO()
    fig.savefig(payload, format="png", dpi=150, bbox_inches="tight", facecolor="white")
    return payload.getvalue()


def _pyplot():
    cache_root = Path(tempfile.gettempdir()) / "janasunani-report-font-cache"
    cache_root.mkdir(parents=True, exist_ok=True)
    os.environ.setdefault("XDG_CACHE_HOME", str(cache_root))
    os.environ.setdefault("MPLCONFIGDIR", str(cache_root / "matplotlib"))
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    return plt


def _runtime_figure(facts: BenchmarkFacts) -> bytes:
    plt = _pyplot()

    text = facts.latency["input_paths"]["text"]["e2e"]
    document = facts.latency["input_paths"]["document"]["e2e"]
    values_minutes = [document["p50"] / 60, text["p50"] / 60, 1.7 * 24 * 60]
    labels = [
        "Janasunani 2.0\nPDF grievance\n(warm p50)",
        "Janasunani 2.0\ntyped grievance\n(warm p50)",
        "Manual routing\n(median, forwarded\ncomplaints)",
    ]
    colors = [FIG_PRIMARY, FIG_BLUE, FIG_ACCENT]
    annotations = [
        f"{document['p50']:.2f} sec\n(n={document['n']})",
        f"{text['p50']:.3f} sec\n(n={text['n']})",
        "1.7 days\n(administrative median)",
    ]
    fig, ax = plt.subplots(figsize=(13.0, 4.7))
    positions = range(len(values_minutes))
    ax.barh(list(positions), values_minutes, color=colors, height=0.52)
    ax.set_xscale("log")
    ax.set_xlim(0.001, 5000)
    ax.set_yticks(list(positions), labels)
    ax.invert_yaxis()
    ax.set_xlabel("Elapsed time (minutes, log scale)", color=FIG_TEXT)
    ax.set_title(
        "Two clocks: technical processing and administrative forwarding",
        loc="left",
        fontsize=18,
        fontweight="bold",
        color=FIG_PRIMARY,
        pad=12,
    )
    ax.grid(axis="x", which="major", linestyle=":", color=FIG_BORDER, linewidth=1.2)
    ax.set_axisbelow(True)
    ax.spines[["top", "right"]].set_visible(False)
    ax.spines[["left", "bottom"]].set_color(FIG_BORDER)
    ax.tick_params(colors=FIG_TEXT_SECONDARY, labelsize=10)
    for position, value, annotation in zip(
        positions, values_minutes, annotations, strict=True
    ):
        ax.annotate(
            annotation,
            xy=(value, position),
            xytext=(8, 0),
            textcoords="offset points",
            va="center",
            ha="left",
            fontsize=10.5,
            color=FIG_PRIMARY,
            bbox={"boxstyle": "round,pad=0.25", "fc": FIG_BACKGROUND, "ec": FIG_BORDER},
        )
    fig.text(
        0.01,
        -0.01,
        f"Technical source: bundle {facts.bundle_id[:12]}, synthetic CPU development run; "
        f"{facts.latency['completed_attempts']}/{facts.latency['attempts']} attempts complete, "
        f"{facts.latency['failed_attempts']} failed. Administrative source: historical forwarded cases. "
        "Technical speed is not officer time saved.",
        fontsize=8.5,
        color=FIG_TEXT_SECONDARY,
    )
    fig.tight_layout(rect=(0, 0.07, 1, 1))
    result = _figure_png(fig)
    plt.close(fig)
    return result


def _pii_figure(facts: BenchmarkFacts) -> bytes:
    plt = _pyplot()

    pii = facts.pii
    by_entity = pii["by_entity"]
    labels = ["PHONE", "AADHAAR", "EMAIL\n(non-gov)", "NAME", "OVERALL"]
    rates = [
        by_entity["PHONE"]["overlap_recall"],
        by_entity["AADHAAR"]["overlap_recall"],
        by_entity["EMAIL"]["overlap_recall"],
        by_entity["NAME"]["overlap_recall"],
        pii["overall"]["overlap_recall"],
    ]
    fig, ax = plt.subplots(figsize=(12.5, 5.1))
    bars = ax.bar(labels, [rate * 100 for rate in rates], color=FIG_BLUE, width=0.54)
    ax.axhline(80.56, color=FIG_ACCENT, linestyle="--", linewidth=2)
    ax.set_ylim(0, 100)
    ax.set_ylabel("Typed overlap recall (%)", color=FIG_TEXT)
    ax.set_title(
        "PII redaction development recall by entity",
        loc="left",
        fontsize=18,
        fontweight="bold",
        color=FIG_PRIMARY,
        pad=12,
    )
    ax.grid(axis="y", linestyle=":", color=FIG_BORDER, linewidth=1.2)
    ax.set_axisbelow(True)
    ax.spines[["top", "right"]].set_visible(False)
    ax.spines[["left", "bottom"]].set_color(FIG_BORDER)
    ax.tick_params(colors=FIG_TEXT_SECONDARY, labelsize=10)
    for bar, rate in zip(bars, rates, strict=True):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 2,
            f"{rate:.1%}",
            ha="center",
            va="bottom",
            fontsize=11,
            fontweight="bold",
            color=FIG_PRIMARY,
        )
    ax.text(
        4.35,
        83,
        "DSI reference 80.56%\n(historical; not a threshold)",
        color=FIG_ACCENT,
        fontsize=9.5,
        bbox={"boxstyle": "round,pad=0.25", "fc": FIG_CARD, "ec": FIG_BORDER},
    )
    fig.text(
        0.01,
        0.01,
        f"Tracked scorecard: {pii['overall']['gold']} scored spans; "
        f"{pii['excluded_by_policy']} excluded by policy; all records in unknown language bucket. "
        f"Coverage-overlap recall {pii['coverage']['overlap_recall']:.1%}; precision is not adjudicated.",
        fontsize=8.5,
        color=FIG_TEXT_SECONDARY,
    )
    fig.tight_layout(rect=(0, 0.07, 1, 1))
    result = _figure_png(fig)
    plt.close(fig)
    return result


def _routing_figure(facts: BenchmarkFacts) -> bytes:
    plt = _pyplot()

    overall = facts.routing_all
    informative = facts.routing_informative
    top1 = [overall["accuracy"] * 100, informative["accuracy"] * 100]
    top3 = [
        overall["top_k_accuracy"]["3"] * 100,
        informative["top_k_accuracy"]["3"] * 100,
    ]

    fig, (quality_ax, time_ax) = plt.subplots(
        1,
        2,
        figsize=(13.0, 5.2),
        gridspec_kw={"width_ratios": [1.25, 1]},
    )
    positions = [0, 1]
    width = 0.34
    top1_bars = quality_ax.bar(
        [position - width / 2 for position in positions],
        top1,
        width,
        label="Top 1",
        color=FIG_PRIMARY,
    )
    top3_bars = quality_ax.bar(
        [position + width / 2 for position in positions],
        top3,
        width,
        label="Top 3",
        color=FIG_BLUE,
    )
    quality_ax.set_title(
        "Historical-destination agreement",
        loc="left",
        fontsize=15,
        fontweight="bold",
        color=FIG_PRIMARY,
    )
    quality_ax.set_xticks(
        positions,
        [
            f"All categories\n(n={overall['n']:,})",
            f"Informative\n(n={informative['n']:,})",
        ],
    )
    quality_ax.set_ylim(0, 100)
    quality_ax.set_ylabel("Chronological holdout accuracy (%)", color=FIG_TEXT)
    quality_ax.legend(frameon=False, loc="upper left", ncols=2)
    quality_ax.grid(axis="y", linestyle=":", color=FIG_BORDER, linewidth=1.2)
    quality_ax.set_axisbelow(True)
    for bars in (top1_bars, top3_bars):
        quality_ax.bar_label(
            bars, fmt="%.1f%%", padding=3, color=FIG_PRIMARY, fontsize=10
        )

    route_days = [48, 23, 9]
    route_labels = ["Route 4", "Route 2", "Step 3"]
    route_colors = [FIG_PRIMARY, FIG_BLUE, FIG_ACCENT]
    route_bars = time_ax.barh(route_labels, route_days, color=route_colors, height=0.5)
    time_ax.invert_yaxis()
    time_ax.set_xlim(0, 55)
    time_ax.set_xlabel("Median elapsed days", color=FIG_TEXT)
    time_ax.set_title(
        "Route timing is descriptive",
        loc="left",
        fontsize=15,
        fontweight="bold",
        color=FIG_PRIMARY,
    )
    time_ax.grid(axis="x", linestyle=":", color=FIG_BORDER, linewidth=1.2)
    time_ax.set_axisbelow(True)
    time_ax.bar_label(route_bars, labels=["48 days", "23 days", "9 days"], padding=4)

    for axis in (quality_ax, time_ax):
        axis.spines[["top", "right"]].set_visible(False)
        axis.spines[["left", "bottom"]].set_color(FIG_BORDER)
        axis.tick_params(colors=FIG_TEXT_SECONDARY, labelsize=10)

    fig.text(
        0.01,
        0.01,
        f"Routing source: bundle {facts.bundle_id[:12]}. Agreement is with the later "
        "historical destination, not adjudicated correct authority. Route-day contrasts "
        "are observational and are not estimates of time saved.",
        fontsize=8.5,
        color=FIG_TEXT_SECONDARY,
    )
    fig.tight_layout(rect=(0, 0.08, 1, 1), w_pad=3.5)
    result = _figure_png(fig)
    plt.close(fig)
    return result


def _replace_docx_media(path: Path, replacements: dict[str, bytes]) -> None:
    with tempfile.NamedTemporaryFile(
        prefix=f".{path.name}.media.", suffix=".docx", dir=path.parent, delete=False
    ) as handle:
        rebuilt = Path(handle.name)
    try:
        with (
            zipfile.ZipFile(path, "r") as source,
            zipfile.ZipFile(
                rebuilt, "w", compression=zipfile.ZIP_DEFLATED
            ) as destination,
        ):
            names = set(source.namelist())
            missing = sorted(set(replacements) - names)
            if missing:
                raise RuntimeError(f"report media changed; missing {missing}")
            for info in source.infolist():
                destination.writestr(
                    info,
                    replacements.get(info.filename, source.read(info.filename)),
                )
        os.replace(rebuilt, path)
    except Exception:
        rebuilt.unlink(missing_ok=True)
        raise


def _key_starting(mapping: dict[str, str], prefix: str) -> str:
    matches = [key for key in mapping if key.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"expected one report replacement beginning {prefix!r}")
    return matches[0]


def _benchmark_replacements(
    facts: BenchmarkFacts,
) -> tuple[
    dict[str, str], dict[tuple[int, int, int, int], str], dict[str, tuple[str, ...]]
]:
    paragraphs = dict(PARAGRAPH_REPLACEMENTS)
    cells = dict(CELL_REPLACEMENTS)
    alternates = {key: tuple(value) for key, value in PARAGRAPH_ALTERNATES.items()}

    latency = facts.latency
    overall = latency["stages"]["e2e"]
    text = latency["input_paths"]["text"]["e2e"]
    document = latency["input_paths"]["document"]["e2e"]
    startup = latency["processor_startup_seconds"]
    attempts = int(latency["attempts"])
    completed = int(latency["completed_attempts"])
    failures = int(latency["failed_attempts"])

    figure_key = _key_starting(paragraphs, "Figure 2 — Time from filing")
    routing_figure_key = _key_starting(paragraphs, "Figure 7 — Left: learned crosswalk")
    timing_key = _key_starting(paragraphs, "Timing the officer feels.")
    previous_figure = paragraphs[figure_key]
    previous_routing_figure = paragraphs[routing_figure_key]
    previous_timing = paragraphs[timing_key]
    alternates[figure_key] = (*alternates.get(figure_key, ()), previous_figure)
    alternates[routing_figure_key] = (
        *alternates.get(routing_figure_key, ()),
        previous_routing_figure,
    )
    alternates[timing_key] = (*alternates.get(timing_key, ()), previous_timing)
    paragraphs[figure_key] = (
        "Figure 2 — Two different clocks. The 1.7-day median is the observed "
        "administrative interval to first forward among about 30,000 forwarded "
        f"complaints. In the CPU development harness, warm typed grievances had "
        f"p50 {text['p50']:.3f} s (n={text['n']}) and PDFs had p50 "
        f"{document['p50']:.3f} s (n={document['n']}). Technical processing "
        "speed is not officer time saved; that requires exposure logging and a "
        "controlled rollout."
    )
    paragraphs[timing_key] = (
        "Technical timing from the versioned development bundle. On the identified "
        f"Apple arm64 laptop, model startup took {startup:.3f} s. The sequential "
        f"synthetic run completed {completed}/{attempts} attempts with {failures} "
        f"failures and retained {overall['n']} post-warm-up measurements across "
        f"{overall['n_clusters']} grievances. Overall mean was "
        f"{overall['mean_seconds']:.3f} s (clustered SE "
        f"{overall['se_seconds']:.3f}), p50 {overall['p50']:.3f} s, p90 "
        f"{overall['p90']:.3f} s and p95 {overall['p95']:.3f} s. Typed text "
        f"averaged {text['mean_seconds']:.3f} s; PDFs averaged "
        f"{document['mean_seconds']:.3f} s. This is technical development timing "
        "using DVC model bytes plus an identified local BART snapshot, not an "
        "approved release benchmark or officer time saved."
    )

    cells[(2, 0, 0, 0)] = f"{text['p50']:.2f} s text / {document['p50']:.2f} s PDF"
    cells[(2, 0, 0, 1)] = (
        f"Warm p50 in the tracked CPU development run (text n={text['n']}; PDF "
        f"n={document['n']}). The 1.7-day first-forward interval is a different "
        "administrative clock; workflow time saved is not yet measured."
    )
    cells[(9, 7, 2, 0)] = (
        f"Tracked CPU development harness: {completed}/{attempts} attempts "
        f"completed, {failures} failed; {overall['n']} warm measurements. Overall "
        f"mean {overall['mean_seconds']:.3f} s, p50 {overall['p50']:.3f} s, p90 "
        f"{overall['p90']:.3f} s, p95 {overall['p95']:.3f} s. Text mean/p50 "
        f"{text['mean_seconds']:.3f}/{text['p50']:.3f} s (n={text['n']}); PDF "
        f"{document['mean_seconds']:.3f}/{document['p50']:.3f} s "
        f"(n={document['n']}). Startup {startup:.3f} s."
    )
    cells[(9, 7, 3, 0)] = (
        "These are sequential synthetic development timings, split by input path. "
        "They are not officer handling time, time to first action, or a release-host "
        "service-level claim."
    )

    action = facts.actionability
    weak = facts.weak_labels
    valid_weak = weak["eligible_ticket_labels"]["valid_single_label"]
    max_tv = weak["office_variation"]["max_total_variation"]
    cells[(25, 12, 1, 0)] = (
        f"Validation-selected {action['selected_candidate']} test n={action['n']}: "
        f"{action['accuracy']:.2%} accuracy; {action['confusion']['true_review']}/"
        f"{action['actual_review']} review recall; {action['confusion']['true_review']}/"
        f"{action['flagged']} review precision; {action['confusion']['false_review']}/"
        f"{action['confusion']['true_actionable'] + action['confusion']['false_review']} "
        f"actionable sent to review. Weak labels n={valid_weak:,}; office max TV "
        f"{max_tv:.3f} (pooling gate fails)."
    )
    cells[(21, 1, 1, 0)] = (
        "Advisory low-signal reason, never a rejection. On the tracked "
        f"{action['n']}-case frontier-adjudicated binary development test, the "
        f"validation-selected local candidate caught "
        f"{action['confusion']['true_review']}/{action['actual_review']} complaints "
        f"needing review and sent {action['confusion']['false_review']}/"
        f"{action['confusion']['true_actionable'] + action['confusion']['false_review']} "
        "ordinary complaints to review. Its checksummed binary artifact is "
        "serving-compatible for advisory review, but it does not assign five-class "
        "reasons and is not release-eligible."
    )

    summary = facts.summary
    cells[(21, 4, 1, 0)] = (
        "Local BART development baseline on an enriched redacted typed-text set: "
        f"{summary['critical_fact_recall']['successes']}/"
        f"{summary['critical_fact_recall']['n']} critical facts retained; 0/"
        f"{summary['generated_n']} unsupported or contradictory cases; "
        f"{summary['usable_without_edit_rate']['successes']}/"
        f"{summary['generated_n']} usable without edit; "
        f"{summary['pii_leak_case_rate']['successes']}/"
        f"{summary['generated_n']} residual-PII cases."
    )
    cells[(21, 4, 2, 0)] = (
        "Single frontier judge and viewed enriched test only. All six judge-marked "
        "skip cases received drafts and all four coherent Odia cases were skipped; "
        "post-summary privacy and paired officer validation are required."
    )

    routing_all = facts.routing_all
    routing_info = facts.routing_informative
    category = facts.categorization
    paragraphs[routing_figure_key] = (
        "Figure 7 — Left: the 60.9/67.5/72.8% crosswalk bars are historical "
        "in-sample resubstitution, retained only as upper-bound context. The "
        f"tracked chronological developmental holdout gives "
        f"{routing_all['accuracy']:.2%} top-1 and "
        f"{routing_all['top_k_accuracy']['3']:.2%} top-3 for live "
        f"category+district features (n={routing_all['n']:,}), or "
        f"{routing_info['accuracy']:.2%}/{routing_info['top_k_accuracy']['3']:.2%} "
        "on informative categories. Right: Route 4's 48-day vs Route 2's 23-day "
        "medians are descriptive and may reflect unobserved complexity; a pilot "
        "must measure any time effect."
    )
    cells[(25, 10, 1, 0)] = (
        f"Category+district: {routing_all['accuracy']:.2%} top-1, "
        f"{routing_all['top_k_accuracy']['3']:.2%} top-3 / n={routing_all['n']:,}; "
        f"informative: {routing_info['accuracy']:.2%} / "
        f"{routing_info['top_k_accuracy']['3']:.2%} / n={routing_info['n']:,}"
    )
    cells[(21, 5, 2, 0)] = (
        "Learned means historical destination agreement, not correct authority or "
        f"best outcome. Chronological developmental top-1 is "
        f"{routing_all['accuracy']:.2%} overall and "
        f"{routing_info['accuracy']:.2%} on informative categories; the older "
        "60.9/67.5/72.8% figures are in-sample."
    )
    cells[(24, 3, 1, 0)] = (
        "Category historical-label agreement: "
        f"{category_benchmark_summary(category)}. "
        "Not policy correctness or release evidence. "
        "Routing historical-destination agreement: "
        f"{routing_all['accuracy']:.2%} top-1 / "
        f"{routing_all['top_k_accuracy']['3']:.2%} top-3 overall "
        f"(n={routing_all['n']:,})."
    )
    cells[(24, 5, 1, 0)] = (
        "Historical destination agreement only. The live category+district "
        f"developmental holdout is {routing_all['accuracy']:.2%} top-1 and "
        f"{routing_all['top_k_accuracy']['3']:.2%} top-3; process-time contrasts "
        "remain descriptive."
    )

    pii = facts.pii
    overall_pii = pii["overall"]
    coverage_pii = pii["coverage"]
    entity = pii["by_entity"]
    cells[(25, 5, 1, 0)] = (
        f"PHONE {entity['PHONE']['overlap_recall']:.1%} / AADHAAR "
        f"{entity['AADHAAR']['overlap_recall']:.1%} / EMAIL "
        f"{entity['EMAIL']['overlap_recall']:.1%} / NAME "
        f"{entity['NAME']['overlap_recall']:.1%} / OVERALL "
        f"{overall_pii['overlap_recall']:.1%} / EXACT "
        f"{overall_pii['exact_recall']:.1%} / COVERAGE "
        f"{coverage_pii['overlap_recall']:.1%} / {overall_pii['gold']} spans"
    )
    cells[(24, 2, 1, 0)] = (
        f"Current development scorecard: {overall_pii['gold']} labelled spans; "
        f"typed overlap recall {overall_pii['overlap_recall']:.1%}, exact recall "
        f"{overall_pii['exact_recall']:.1%}, and untyped coverage-overlap recall "
        f"{coverage_pii['overlap_recall']:.1%}. All records are in an unknown "
        "language bucket."
    )
    cells[(9, 3, 2, 0)] = (
        "Presidio rebuild (fully in-process). Tracked development scorecard: "
        f"{overall_pii['gold']} scored spans, {pii['excluded_by_policy']} excluded "
        f"by policy; overlap recall {overall_pii['overlap_recall']:.1%}, exact "
        f"recall {overall_pii['exact_recall']:.1%}, coverage-overlap "
        f"{coverage_pii['overlap_recall']:.1%}. PHONE "
        f"{entity['PHONE']['overlap_recall']:.1%}; AADHAAR "
        f"{entity['AADHAAR']['overlap_recall']:.1%}; EMAIL "
        f"{entity['EMAIL']['overlap_recall']:.1%}; NAME "
        f"{entity['NAME']['overlap_recall']:.1%}. All records are in the unknown "
        "language bucket; precision is not adjudicated."
    )

    cells[(25, 18, 0, 0)] = "Runtime — tracked CPU development bundle"
    cells[(25, 18, 1, 0)] = (
        f"overall mean {overall['mean_seconds']:.3f} s / p50 {overall['p50']:.3f} / "
        f"p90 {overall['p90']:.3f} / p95 {overall['p95']:.3f}; text mean "
        f"{text['mean_seconds']:.3f} s (n={text['n']}); PDF mean "
        f"{document['mean_seconds']:.3f} s (n={document['n']}); "
        f"{completed}/{attempts} complete, {failures} failed"
    )
    cells[(25, 18, 2, 0)] = (
        f"outputs/benchmark/full_benchmark.json; bundle {facts.bundle_id}"
    )
    cells[(25, 18, 3, 0)] = (
        "dvc pull outputs/benchmark/latency.json.dvc && "
        "dvc repro --single-item full-benchmark-bundle"
    )
    cells[(25, 21, 0, 0)] = "Impact — operational and citizen"
    cells[(25, 21, 1, 0)] = (
        f"not measured: {facts.impact_available_required}/{facts.impact_required} "
        "required impact artifacts available; no exposure/outcome or satisfaction "
        "effect estimate"
    )
    cells[(25, 21, 2, 0)] = (
        f"outputs/benchmark/full_benchmark.json; publication_ready="
        f"{str(facts.publication_ready).lower()}"
    )
    cells[(25, 21, 3, 0)] = (
        "Implement exposure, officer-decision, resolution and fixed-horizon "
        "satisfaction events; then run the locked pilot estimator"
    )
    return paragraphs, cells, alternates


def patch_report(source: Path, destination: Path, *, benchmark_bundle: Path) -> None:
    facts = load_benchmark_facts(benchmark_bundle)
    paragraph_replacements, cell_replacements, paragraph_alternates = (
        _benchmark_replacements(facts)
    )
    document = Document(source)
    current = {paragraph.text: paragraph for paragraph in document.paragraphs}

    def dynamic_source(old: str):
        prefixes: tuple[str, ...] = ()
        if old.startswith("Figure 2 — Time from filing"):
            prefixes = ("Figure 2 — Two different clocks.",)
        elif old.startswith("Timing the officer feels."):
            prefixes = ("Technical timing from the versioned development bundle.",)
        return next(
            (
                paragraph
                for text, paragraph in current.items()
                if any(text.startswith(prefix) for prefix in prefixes)
            ),
            None,
        )

    missing = sorted(
        old
        for old, new in paragraph_replacements.items()
        if old not in current
        and new not in current
        and not any(
            alternate in current for alternate in paragraph_alternates.get(old, ())
        )
        and dynamic_source(old) is None
    )
    if missing:
        preview = "\n".join(f"- {text[:180]}" for text in missing)
        raise RuntimeError(
            f"report paragraphs changed; missing {len(missing)} expected texts:\n{preview}"
        )
    for old, new in paragraph_replacements.items():
        if old in current:
            _set_paragraph(current[old], new)
            if new.startswith("Technical timing from the versioned development bundle"):
                current[old].paragraph_format.space_before = Pt(7)
            continue
        for alternate in paragraph_alternates.get(old, ()):
            if alternate in current:
                _set_paragraph(current[alternate], new)
                if new.startswith(
                    "Technical timing from the versioned development bundle"
                ):
                    current[alternate].paragraph_format.space_before = Pt(7)
                break
        else:
            paragraph = dynamic_source(old)
            if paragraph is not None:
                _set_paragraph(paragraph, new)

    for (
        table_index,
        row_index,
        cell_index,
        paragraph_index,
    ), text in cell_replacements.items():
        cell = document.tables[table_index].cell(row_index, cell_index)
        if paragraph_index >= len(cell.paragraphs):
            raise RuntimeError(
                f"missing paragraph {paragraph_index} in table {table_index} "
                f"row {row_index} cell {cell_index}"
            )
        _set_paragraph(cell.paragraphs[paragraph_index], text)

    for table in document.tables:
        if table.rows:
            _set_repeat_table_header(table.rows[0])

    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        prefix=f".{destination.name}.",
        suffix=".docx",
        dir=destination.parent,
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        _replace_docx_media(
            temporary,
            {
                "word/media/image1.png": _runtime_figure(facts),
                "word/media/image2.png": _pii_figure(facts),
                "word/media/image6.png": _routing_figure(facts),
            },
        )
        canonicalize_docx_archive(temporary)
        os.replace(temporary, destination)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_REPORT)
    parser.add_argument("--benchmark-bundle", type=Path, default=DEFAULT_BUNDLE)
    args = parser.parse_args()
    patch_report(
        args.input,
        args.output,
        benchmark_bundle=args.benchmark_bundle,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
