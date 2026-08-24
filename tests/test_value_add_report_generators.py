from pathlib import Path
import sys
import zipfile

from docx import Document

from janasunani.evaluation.value_add_benchmark_facts import BenchmarkFacts

SCRIPTS = Path(__file__).parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

import create_officer_brief as officer  # noqa: E402
import create_public_systems_capability_brief as public  # noqa: E402
from docx_archive import CANONICAL_ZIP_TIMESTAMP  # noqa: E402
import update_value_add_report as report  # noqa: E402


def benchmark_facts() -> BenchmarkFacts:
    latency_stat = {
        "n": 4,
        "n_clusters": 2,
        "mean_seconds": 2.5,
        "se_seconds": 0.1,
        "p50": 2.3,
        "p90": 3.0,
        "p95": 3.2,
    }
    return BenchmarkFacts(
        bundle_id="test-bundle-1234567890",
        publication_ready=False,
        latency={
            "input_paths": {
                "text": {"e2e": latency_stat},
                "document": {"e2e": {**latency_stat, "p50": 8.5}},
            },
            "stages": {"e2e": latency_stat},
            "processor_startup_seconds": 1.2,
            "attempts": 8,
            "completed_attempts": 8,
            "failed_attempts": 0,
        },
        actionability={
            "selected_candidate": "tfidf",
            "release_eligible": False,
            "n": 50,
            "accuracy": 0.82,
            "actual_review": 20,
            "flagged": 23,
            "confusion": {
                "true_review": 18,
                "false_review": 5,
                "true_actionable": 30,
            },
        },
        categorization={
            "accuracy": 0.4514,
            "top_k_accuracy": {"3": 0.6904},
            "macro_f1": 0.42,
            "n": 208267,
        },
        weak_labels={
            "eligible_ticket_labels": {"valid_single_label": 106683},
            "office_variation": {"max_total_variation": 0.522},
        },
        pii={
            "overall": {"gold": 50, "overlap_recall": 0.84, "exact_recall": 0.7},
            "coverage": {"overlap_recall": 0.86},
            "by_entity": {
                entity: {"overlap_recall": value}
                for entity, value in {
                    "PHONE": 0.9,
                    "AADHAAR": 0.8,
                    "EMAIL": 0.75,
                    "NAME": 0.7,
                }.items()
            },
            "excluded_by_policy": 2,
        },
        routing_all={
            "accuracy": 0.4514,
            "top_k_accuracy": {"3": 0.6904},
            "n": 208267,
        },
        routing_informative={
            "accuracy": 0.5496,
            "top_k_accuracy": {"3": 0.7968},
            "n": 120000,
        },
        summary={
            "critical_fact_recall": {"successes": 55, "n": 84},
            "usable_without_edit_rate": {"successes": 8},
            "pii_leak_case_rate": {"successes": 4},
            "generated_n": 26,
        },
        section_status={
            "speed": {"available_required": 1, "required": 1},
            "accuracy": {"available_required": 6, "required": 6},
            "impact": {"available_required": 0, "required": 2},
        },
    )


def _document_text(path: Path) -> str:
    document = Document(path)
    text = [paragraph.text for paragraph in document.paragraphs]
    text.extend(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    return "\n".join(text)


def test_actual_word_generators_create_bundle_backed_outputs(tmp_path, monkeypatch):
    facts = benchmark_facts()
    for module in (officer, public, report):
        monkeypatch.setattr(module, "load_benchmark_facts", lambda _path: facts)

    officer_path = tmp_path / "officer.docx"
    public_path = tmp_path / "public.docx"
    report_path = tmp_path / "report.docx"
    officer.create_brief(officer_path, benchmark_bundle=Path("unused.json"))
    public.create_brief(public_path, benchmark_bundle=Path("unused.json"))
    report.patch_report(
        report.DEFAULT_SOURCE,
        report_path,
        benchmark_bundle=Path("unused.json"),
    )

    for output in (officer_path, public_path, report_path):
        assert output.stat().st_size > 10_000
        assert "test-bundle-1234567890" in _document_text(output)

    document = Document(report_path)
    assert len(document.inline_shapes) > 0


def test_word_generators_produce_reproducible_archives(tmp_path, monkeypatch):
    facts = benchmark_facts()
    for module in (officer, public, report):
        monkeypatch.setattr(module, "load_benchmark_facts", lambda _path: facts)

    first = tmp_path / "first"
    second = tmp_path / "second"
    first.mkdir()
    second.mkdir()

    outputs = []
    for directory in (first, second):
        officer_path = directory / "officer.docx"
        public_path = directory / "public.docx"
        report_path = directory / "report.docx"
        officer.create_brief(officer_path, benchmark_bundle=Path("unused.json"))
        public.create_brief(public_path, benchmark_bundle=Path("unused.json"))
        report.patch_report(
            report.DEFAULT_SOURCE,
            report_path,
            benchmark_bundle=Path("unused.json"),
        )
        outputs.append((officer_path, public_path, report_path))

    for first_output, second_output in zip(*outputs, strict=True):
        assert first_output.read_bytes() == second_output.read_bytes()
        with zipfile.ZipFile(first_output) as archive:
            assert {member.date_time for member in archive.infolist()} == {
                CANONICAL_ZIP_TIMESTAMP
            }
